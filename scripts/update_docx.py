import docx

def delete_paragraph(paragraph):
    p = paragraph._element
    p.getparent().remove(p)
    p._element = p._parent = None

def main():
    doc_path = '/home/truongan/Downloads/resume_truongan.docx'
    doc = docx.Document(doc_path)
    
    # 1. Cập nhật Education (Paragraph 5)
    # B.Sc. Computer Science  |  Vietnam National University (VNU) Sep 2022 – Jun 2026
    for p in doc.paragraphs:
        if "B.Sc. Computer Science" in p.text and "Vietnam National University" in p.text:
            print("Cập nhật Education:", p.text)
            p.text = "B.Sc. Computer and Information Science  |  University of Science - VNU   Sep 2022 – Jun 2026"
            # Giữ nguyên style bold cho phần text
            for run in p.runs:
                if "B.Sc. Computer" in run.text:
                    run.text = "B.Sc. Computer and Information Science  |  University of Science - VNU"
    
    # 2. Tìm và xóa dự án Unity Game và Line Follower
    # Chúng ta sẽ duyệt ngược từ dưới lên để tránh lỗi lệch index khi xóa
    paragraphs_to_remove = []
    
    # Lưu danh sách chỉ số các đoạn cần xóa
    for i, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        # Xóa Su Thi Viet Nam
        if "Su Thi Viet Nam" in text or "Developed core gameplay architecture in Unity" in text:
            paragraphs_to_remove.append(p)
        # Xóa Autonomous Line-Follower
        if "Autonomous Line-Follower" in text or "Designed physical electrical layout with IR" in text or "Developed and fine-tuned a closed-loop PID" in text:
            paragraphs_to_remove.append(p)

    # Tiến hành xóa các đoạn được đánh dấu
    for p in paragraphs_to_remove:
        print("Xóa đoạn văn bản:", p.text)
        try:
            delete_paragraph(p)
        except Exception as e:
            print(f"Lỗi khi xóa: {e}")
            
    # Lưu lại file
    doc.save(doc_path)
    print("Đã cập nhật và lưu file thành công!")

if __name__ == "__main__":
    main()
